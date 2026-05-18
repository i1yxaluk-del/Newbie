"""
Общий модуль форматирования для пакета договоров MSPShield.

Создаёт единый визуальный стиль для всех .docx-файлов: шрифт, отступы,
заголовки, плейсхолдеры. Не содержит юридического контента —
только утилиты.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Константы стиля
# ---------------------------------------------------------------------------

FONT_BODY = "PT Serif"
FONT_HEAD = "PT Sans"
FONT_FALLBACK_BODY = "Times New Roman"
FONT_FALLBACK_HEAD = "Arial"

INK = RGBColor(0x1A, 0x1A, 0x1A)
INK_MUTED = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0x3E, 0x5C, 0x3A)

HIGHLIGHT_PLACEHOLDER = "yellow"

PAGE_MARGIN_TOP = Cm(2.0)
PAGE_MARGIN_BOTTOM = Cm(2.0)
PAGE_MARGIN_LEFT = Cm(2.5)
PAGE_MARGIN_RIGHT = Cm(2.0)


# ---------------------------------------------------------------------------
# Базовые билдеры
# ---------------------------------------------------------------------------


def new_document() -> Document:
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = PAGE_MARGIN_TOP
        section.bottom_margin = PAGE_MARGIN_BOTTOM
        section.left_margin = PAGE_MARGIN_LEFT
        section.right_margin = PAGE_MARGIN_RIGHT

    # Базовый стиль "Normal" — это шрифт по умолчанию для всего документа.
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), FONT_BODY)
    rfonts.set(qn("w:hAnsi"), FONT_BODY)
    rfonts.set(qn("w:cs"), FONT_FALLBACK_BODY)
    rfonts.set(qn("w:eastAsia"), FONT_FALLBACK_BODY)

    return doc


def add_title(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    for line in text.split("\n"):
        run = p.add_run(line)
        run.font.name = FONT_HEAD
        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = INK
        # Принудительный перенос строки в одном параграфе.
        if line != text.split("\n")[-1]:
            run.add_break()


def add_subtitle(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.font.name = FONT_BODY
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = INK_MUTED


def add_h1(doc, text: str, *, before: int = 18, after: int = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT_HEAD
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = INK


def add_h2(doc, text: str, *, before: int = 10, after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT_HEAD
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = INK


def _apply_body_font(run, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT_BODY
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    run.bold = bold
    run.italic = italic


def add_clause(doc, number: str, text: str) -> None:
    """Добавить пункт договора (например, 2.1.3).

    Текст может содержать токены вида **жирный** и __плейсхолдер__:
    - **...** — выделение жирным
    - __...__ — плейсхолдер с жёлтой подсветкой для заполнения
    """
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Номер пункта — отрицательный отступ, чтобы создать "висячий" эффект.
    p.paragraph_format.first_line_indent = Cm(-1.0)
    num_run = p.add_run(f"{number}. ")
    _apply_body_font(num_run, bold=True)

    _add_rich_text(p, text)


def add_paragraph(doc, text: str, *, indent: float = 0.0, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _add_rich_text(p, text, italic=italic)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(2)
    _add_rich_text(p, text)


def add_note(doc, text: str) -> None:
    """Информационный блок (курсив, отступ, серый цвет)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_body_font(run, italic=True)
    run.font.color.rgb = INK_MUTED


def _add_rich_text(p, text: str, *, italic: bool = False) -> None:
    """Парсит текст на жирные участки (**...**) и плейсхолдеры (__...__)."""
    i = 0
    n = len(text)
    while i < n:
        if text[i:i + 2] == "**":
            end = text.find("**", i + 2)
            if end == -1:
                run = p.add_run(text[i:])
                _apply_body_font(run, italic=italic)
                break
            run = p.add_run(text[i + 2:end])
            _apply_body_font(run, bold=True, italic=italic)
            i = end + 2
        elif text[i:i + 2] == "__":
            end = text.find("__", i + 2)
            if end == -1:
                run = p.add_run(text[i:])
                _apply_body_font(run, italic=italic)
                break
            run = p.add_run(text[i + 2:end])
            _apply_body_font(run, italic=italic)
            _highlight_run(run, HIGHLIGHT_PLACEHOLDER)
            i = end + 2
        else:
            # Найти ближайший токен.
            next_b = text.find("**", i)
            next_p = text.find("__", i)
            candidates = [x for x in (next_b, next_p) if x != -1]
            stop = min(candidates) if candidates else n
            run = p.add_run(text[i:stop])
            _apply_body_font(run, italic=italic)
            i = stop


def _highlight_run(run, color: str) -> None:
    rpr = run._element.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), color)
    rpr.append(highlight)


def add_table(
    doc,
    headers: list[str],
    rows: list[list[str]],
    *,
    col_widths_cm: list[float] | None = None,
) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        run.font.name = FONT_HEAD
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(cell, "2D2D2D")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row in enumerate(rows):
        for ci, value in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            _add_rich_text(p, value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if ri % 2 == 1:
                _shade_cell(cell, "F5F2EC")

    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_signature_block(doc, parties: list[tuple[str, list[str]]]) -> None:
    """Добавить блок подписей в виде таблицы 1xN."""
    table = doc.add_table(rows=1, cols=len(parties))
    table.autofit = True

    for i, (title, lines) in enumerate(parties):
        cell = table.rows[0].cells[i]
        cell.text = ""
        h = cell.paragraphs[0]
        run = h.add_run(title)
        run.font.name = FONT_HEAD
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = INK

        for line in lines:
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            _add_rich_text(p, line)

        # Подписная линия.
        sig_p = cell.add_paragraph()
        sig_p.paragraph_format.space_before = Pt(16)
        sig_p.paragraph_format.space_after = Pt(2)
        run = sig_p.add_run("_____________________________")
        _apply_body_font(run)
        label_p = cell.add_paragraph()
        run = label_p.add_run("подпись / дата / печать (при наличии)")
        _apply_body_font(run, italic=True)
        run.font.size = Pt(9)
        run.font.color.rgb = INK_MUTED


def add_page_break(doc) -> None:
    doc.add_page_break()


def save(doc, filename: str) -> Path:
    out_dir = Path(__file__).parent / "build"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / filename
    doc.save(str(path))
    return path
